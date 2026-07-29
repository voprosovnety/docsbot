"""Turn an uploaded file into embeddable chunks.

Two steps, kept separate so each is testable on its own:
  extract_text() — bytes to (text, page) pairs
  chunk_text()   — long text to overlapping windows that fit a prompt
"""

from __future__ import annotations

import io
import re
from pathlib import Path

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt", ".markdown"}


class UnsupportedFormat(Exception):
    """Raised for a file extension the bot cannot read."""


class EmptyDocument(Exception):
    """Raised when a file parses fine but contains no extractable text."""


class UnreadableDocument(Exception):
    """Raised when a file claims a supported format but cannot be parsed."""


def extract_text(filename: str, data: bytes) -> list[tuple[str, int | None]]:
    """Return (text, page_number) pairs. `page_number` is None for flat formats."""
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFormat(suffix or filename)

    if suffix == ".pdf":
        pages = _extract_pdf(data)
    elif suffix == ".docx":
        pages = [(_extract_docx(data), None)]
    else:
        pages = [(data.decode("utf-8", errors="replace"), None)]

    cleaned = [(_normalise(text), page) for text, page in pages]
    non_empty = [(text, page) for text, page in cleaned if text]
    if not non_empty:
        raise EmptyDocument(filename)
    return non_empty


def _extract_pdf(data: bytes) -> list[tuple[str, int | None]]:
    from pypdf import PdfReader
    from pypdf.errors import PyPdfError

    try:
        reader = PdfReader(io.BytesIO(data))
        # Pages are 1-indexed in the citation, matching what a reader sees.
        return [
            (page.extract_text() or "", number)
            for number, page in enumerate(reader.pages, 1)
        ]
    # PyPdfError is the base class: EmptyFileError, PdfReadError,
    # WrongPasswordError and friends all inherit from it.
    except PyPdfError as exc:
        # Truncated uploads, password-protected files, and anything that is not
        # really a PDF land here. Without this they'd surface as a generic crash.
        raise UnreadableDocument(str(exc)) from exc


def _extract_docx(data: bytes) -> str:
    import zipfile

    import docx
    from docx.opc.exceptions import PackageNotFoundError

    try:
        document = docx.Document(io.BytesIO(data))
    # A .docx is a zip archive, so a file that isn't one surfaces as BadZipFile
    # rather than a python-docx error of its own.
    except (PackageNotFoundError, zipfile.BadZipFile, KeyError, ValueError) as exc:
        raise UnreadableDocument(str(exc)) from exc
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _normalise(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def chunk_text(
    text: str, *, chunk_size: int, overlap: int, page: int | None = None
) -> list[tuple[str, int | None]]:
    """Split text into overlapping windows, preferring paragraph boundaries.

    Splitting mid-sentence costs retrieval quality, so each window is trimmed
    back to the last blank line or sentence end when one is available in the
    final quarter of the window.
    """
    if overlap >= chunk_size:
        raise ValueError("overlap must be smaller than chunk_size")

    text = text.strip()
    if not text:
        return []
    if len(text) <= chunk_size:
        return [(text, page)]

    chunks: list[tuple[str, int | None]] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        if end < len(text):
            end = _find_boundary(text, start, end, chunk_size)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append((chunk, page))
        if end >= len(text):
            break
        start = max(end - overlap, start + 1)
    return chunks


def _find_boundary(text: str, start: int, end: int, chunk_size: int) -> int:
    """Pull `end` back to a natural break, if one sits in the last quarter."""
    window_floor = start + (chunk_size * 3 // 4)
    for separator in ("\n\n", ". ", "\n"):
        found = text.rfind(separator, window_floor, end)
        if found != -1:
            return found + len(separator)
    return end


def build_chunks(
    pages: list[tuple[str, int | None]], *, chunk_size: int, overlap: int
) -> list[tuple[str, int | None]]:
    """Chunk every page and flatten the result, keeping page numbers attached."""
    chunks: list[tuple[str, int | None]] = []
    for text, page in pages:
        chunks.extend(chunk_text(text, chunk_size=chunk_size, overlap=overlap, page=page))
    return chunks
